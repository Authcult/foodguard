"""
生成面试用项目详解 Word 文档
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# ---- 样式设置 ----
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ---- 标题 ----
title = doc.add_heading('食鉴（FoodGuard）项目详解', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('食品配料智能分析 Agent — 基于 RAG + LangGraph 的个人学习项目')
doc.add_paragraph('')

# ============================================================
# 第一章
# ============================================================
doc.add_heading('一、项目概述', level=1)
doc.add_paragraph(
    '"食鉴"是一个食品配料智能分析 Agent。用户输入食品配料表（如拍照或手动输入），'
    '系统用自然语言对话式地帮用户解读每种配料的作用、标注风险等级、检测过敏原、对比不同食品。'
)
doc.add_paragraph(
    '本项目是一个学习型项目，目标是通过实践掌握 RAG（检索增强生成）+ LangChain + LangGraph 的完整技术栈。'
    '适合在实习面试中展示 LLM 应用工程落地能力。'
)

# ============================================================
# 第二章
# ============================================================
doc.add_heading('二、核心功能', level=1)
functions = [
    ('配料解读', '逐条解释配料含义、作用、安全性，用通俗语言让消费者理解食品标签'),
    ('风险标注', '对每个配料标注 🟢安全/🟡注意/🔴回避 三级风险，并说明原因'),
    ('过敏原检测', '根据用户保存的过敏史，自动检测食品中是否含有危险成分，包括直接匹配、别名匹配和交叉反应'),
    ('对比分析', '两款同类食品配料对比分析，推荐更健康的选择'),
    ('多轮对话', '支持追问和深入讨论，Agent 自动识别用户意图并路由到对应功能'),
]
for name, desc in functions:
    p = doc.add_paragraph()
    run = p.add_run(f'🔹 {name}：')
    run.bold = True
    p.add_run(desc)

# ============================================================
# 第三章
# ============================================================
doc.add_heading('三、技术架构', level=1)
doc.add_paragraph(
    '系统采用分层架构，从上到下依次为：前端层 → Agent 编排层 → RAG Chain 层 → 数据存储层。'
    '各层职责清晰，层与层之间通过接口解耦。'
)

# 架构图
doc.add_heading('3.1 架构全景图', level=2)
arch_text = """
┌─────────────────────────────────────────────────┐
│                  Streamlit 前端                   │
│          chat_input / chat_message / 侧边栏       │
└──────────────────────┬──────────────────────────┘
                       │
┌──────────────────────▼──────────────────────────┐
│              LangGraph Agent 编排层                │
│  ┌─────────┐    ┌──────────────────────────┐    │
│  │ Router  │───▶│ interpret / risk /        │    │
│  │ 意图识别 │    │ compare / allergy 节点    │    │
│  └─────────┘    └──────────────────────────┘    │
└──────────────────────┬──────────────────────────┘
                       │
┌──────────────────────▼──────────────────────────┐
│               LangChain RAG 链路层                │
│   ① 检索 (ChromaDB + BGE-M3 Embedding)           │
│   ② 增强 (Prompt 模板填入检索结果)                │
│   ③ 生成 (DeepSeek LLM 生成回答)                  │
└──────────────────────┬──────────────────────────┘
                       │
┌──────────────────────▼──────────────────────────┐
│                  数据存储层                        │
│   ChromaDB 向量库  │  JSON 知识库 (145 条)         │
│   BGE-M3 模型      │  用户画像 (过敏史)            │
└─────────────────────────────────────────────────┘
"""
p = doc.add_paragraph()
run = p.add_run(arch_text)
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_heading('3.2 数据流', level=2)
doc.add_paragraph(
    '用户输入 → Router 识别意图 → 选择对应 Chain → Retriever 向量检索 → '
    '拼接 Prompt 模板 → DeepSeek LLM 生成回答 → 返回用户'
)

# ============================================================
# 第四章
# ============================================================
doc.add_heading('四、各层详解', level=1)

# 4.1
doc.add_heading('4.1 数据层', level=2)
doc.add_paragraph(
    '知识库是 RAG 系统的根基。本项目知识库包含 145 条食品添加剂数据，来自两个渠道：'
)
doc.add_paragraph(
    '手工标注（88 条）：人工审核过的核心添加剂数据，包含风险等级（safe/caution/avoid）、'
    '风险原因说明、儿童/孕妇安全性、过敏原信息、每日摄入限量等。这是系统最有价值的数据。',
    style='List Bullet'
)
doc.add_paragraph(
    'GB2760-2024 爬虫（57 条补充）：从食品伙伴网（2760.foodmate.net）爬取的标准数据，'
    '包含 CNS/INS 编码和功能分类。由于未经人工审核，标记为"待审核"。',
    style='List Bullet'
)
doc.add_paragraph(
    '数据处理流程：爬虫脚本（crawl_gb2760.py）遍历 22 个功能分类页面，解析 HTML 表格，'
    '提取添加剂名称、CNS、INS、功能。然后 process_data.py 以人工标注为主，爬虫数据为辅，'
    '按名称去重合并。每条添加剂构造成 LangChain Document 对象，page_content 用于 Embedding，'
    'metadata 保留结构化字段。'
)

# 4.2
doc.add_heading('4.2 向量检索层', level=2)
doc.add_paragraph(
    'Embedding 模型选用 BGE-M3（BAAI/bge-m3），这是目前中文检索效果最好的开源模型之一。'
    '它将文本转换为 1024 维向量，语义相近的文本在向量空间中距离更近。国内通过 ModelScope '
    '(modelscope.cn) 下载，避免 HuggingFace 被墙的问题。'
)
doc.add_paragraph(
    '向量数据库选用 ChromaDB，本地持久化到 data/chroma_db/ 目录，无需额外部署服务。'
    '每条添加剂的 page_content 被 Embedding 后在 ChromaDB 中建立索引。'
    '检索时，用户问题同样经过 BGE-M3 Embedding，与 145 条知识做余弦相似度计算，返回 Top K 最相关条目。'
)
doc.add_paragraph(
    '关键设计：page_content 和 metadata 分离。page_content 包含尽可能丰富的描述文本（名称、别名、功能、风险说明、使用范围），'
    '用于提升检索命中率。metadata 保留结构化字段（risk_level、children_safe 等），供后续 Chain 直接取值，避免二次查库。'
)

# 4.3
doc.add_heading('4.3 Chain 层（RAG 核心）', level=2)
doc.add_paragraph(
    'Chain 是 RAG 的执行单元。每个 Chain 遵循相同模式：检索 → 组装 Prompt → 调用 LLM → 输出文本。'
    '使用 LangChain 的 LCEL（LangChain Expression Language）管道语法串联：'
)
chain_code = '''
chain = (
    {
        "context": lambda x: format_docs(retriever.invoke(x["question"])),
        "question": lambda x: x["question"],
    }
    | prompt              # 将检索结果填入 Prompt 模板
    | llm                 # 调用 DeepSeek 生成回答
    | StrOutputParser()   # 提取纯文本输出
)
'''
p = doc.add_paragraph()
run = p.add_run(chain_code)
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_paragraph(
    '系统包含四个 Chain：'
)
doc.add_paragraph('interpret_chain：配料解读，用通俗语言解释配料作用和安全性', style='List Bullet')
doc.add_paragraph('risk_chain：风险标注，按 🟢🟡🔴 等级标注每个配料', style='List Bullet')
doc.add_paragraph('compare_chain：对比分析，从多维度比较两款食品', style='List Bullet')
doc.add_paragraph('allergy_chain：过敏检测，根据用户过敏史排查风险', style='List Bullet')

# 4.4
doc.add_heading('4.4 Agent 编排层（LangGraph）', level=2)
doc.add_paragraph(
    'Agent 层负责根据用户意图自动选择对应的 Chain。使用 LangGraph 的 StateGraph 实现。'
)
doc.add_paragraph(
    '意图路由（Router）：优先使用关键词匹配（零成本、毫秒级响应），覆盖 90% 以上场景。'
    '当关键词无匹配时，自动降级为 LLM 路由作为兜底。这是一个工程权衡——优先规则引擎，LLM 作为安全网。'
)
doc.add_paragraph(
    '五个意图：interpret（解读）、risk（风险）、compare（对比）、allergy（过敏）、general（一般对话）。'
)
doc.add_paragraph(
    'StateGraph 将 Router 和四个 Chain 编排为条件分支图：用户输入 → router_node → '
    '根据意图 → interpret_node / risk_node / compare_node / allergy_node / general_node → 结束。'
    '每个节点独立运行，新增功能只需添加节点和边，不影响现有功能。'
)

# 4.5
doc.add_heading('4.5 前端层（Streamlit）', level=2)
doc.add_paragraph(
    'Streamlit 提供纯 Python 的 Web UI。使用 chat_message 和 chat_input 组件实现对话式界面，'
    '侧边栏提供过敏史设置。使用 @st.cache_resource 确保 LLM 和 Embedding 模型只初始化一次，'
    '后续请求复用。'
)

# ============================================================
# 第五章
# ============================================================
doc.add_heading('五、技术栈选型理由', level=1)

tech_reasons = [
    ('DeepSeek API',
     '性价比极高的国产大模型，兼容 OpenAI 接口（通过 langchain-openai 的 ChatOpenAI 调用），'
     '无需修改代码即可切换为其他兼容模型（如 Qwen）。'),
    ('LangChain',
     '业界最成熟的 LLM 应用框架。LCEL 管道语法让 RAG 流程清晰可读，'
     '内置 Retriever、Prompt Template、Output Parser 等组件，不需要造轮子。'),
    ('LangGraph',
     'LangChain 生态的 Agent 编排工具。StateGraph 提供节点、边、条件边等抽象，'
     '适合构建有状态的多步骤 Agent。相比手写 if-else 路由，扩展性更好。'),
    ('ChromaDB',
     '轻量级本地向量数据库，零部署。适合学习和原型开发。如需生产化可平滑升级到 Milvus 或 Pinecone。'),
    ('BGE-M3',
     '中文 Embedding 效果最好的开源模型之一。多语言支持、8192 token 最大长度、'
     '1024 维输出。通过 ModelScope 国内可下载。'),
    ('Streamlit',
     '最快捷的 LLM 应用前端方案。原生支持聊天组件，无需写 HTML/CSS/JS。'
     '适合快速验证和原型展示。'),
]
for tech, reason in tech_reasons:
    p = doc.add_paragraph()
    run = p.add_run(f'{tech}：')
    run.bold = True
    p.add_run(reason)

# ============================================================
# 第六章
# ============================================================
doc.add_heading('六、面试问答准备', level=1)

qa_pairs = [
    ('Q: 这个项目是做什么的？',
     'A: 一个基于 RAG + LangGraph 的食品配料智能分析 Agent。用户输入食品配料表，系统能自动解读每种配料的作用、'
     '标注风险等级（安全/注意/回避）、检测过敏原、对比不同食品。技术栈包括 LangChain、LangGraph、'
     'ChromaDB、BGE-M3 Embedding、DeepSeek LLM、Streamlit。我从数据采集到前端全部独立完成。'),

    ('Q: RAG 在你的项目中具体是怎么工作的？',
     'A: 分三步。第一步检索：用户问题通过 BGE-M3 Embedding 转为向量，在 ChromaDB 中做余弦相似度检索，'
     '返回 Top 10 最相关知识条目。第二步增强：将检索结果格式化后填入 Prompt 模板的 {context} 占位符，'
     '与用户原始问题拼接成完整 Prompt。第三步生成：完整 Prompt 发送给 DeepSeek，LLM 基于提供的信息生成回答。'
     'RAG 的核心价值是让 LLM 回答有据可查，不是凭记忆胡编。'),

    ('Q: 为什么用 LangGraph 而不是直接调 Chain？',
     'A: 用户的问题类型多样——解读配料、查风险、比产品、测过敏。如果只有一个 Chain，所有问题都用同一个 Prompt 模板，'
     '效果很差。LangGraph 的 StateGraph 让我用 Router 自动识别意图，分发到四个专用 Chain。'
     '每个 Chain 有自己优化的 Prompt 模板和参数。而且新增功能时只需加节点和边，不影响现有功能。'),

    ('Q: 为什么用 ChromaDB 而不是 FAISS 或其他向量数据库？',
     'A: ChromaDB 开箱即用、本地持久化，不需要额外部署服务，适合个人项目和学习场景。'
     '和 LangChain 集成也很成熟。FAISS 更轻量但不支持持久化，Milvus/Pinecone 太重适合生产环境。'),

    ('Q: 如何处理 Embedding 模型的下载问题（国内网络限制）？',
     'A: 优先从 ModelScope（阿里维护的国内模型平台）下载 BGE-M3，如果 ModelScope 不可用则自动尝试 HuggingFace 镜像。'
     '下载后模型缓存到本地，后续直接加载，不需要重复下载。'),

    ('Q: 遇到的最大挑战是什么？',
     'A: 两个。一是数据质量——GB2760 官网只提供原始标准数据，没有安全性评估。我手动标注了 88 种常见添加剂的风险等级、'
     '特殊人群提醒等。二是意图路由——用户表达方式千差万别，关键词路由只能覆盖 90% 场景，我加了 LLM 路由作为兜底。'),

    ('Q: 如果要把这个项目产品化，你会怎么改进？',
     'A: 第一，知识库扩充——覆盖 GB2760 全部 2000+ 添加剂，引入科研论文作为 RAG 知识源。'
     '第二，支持 OCR——用户拍照上传配料表，自动识别并解析。'
     '第三，模型上生产——换用 Milvus 做向量库，部署到服务器，加用户认证和对话记录。'
     '第四，路由升级——用微调的小模型做意图分类，速度和准确性兼顾。'),
]
for q, a in qa_pairs:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    p = doc.add_paragraph(a)
    p.paragraph_format.space_after = Pt(12)

# ============================================================
# 第七章
# ============================================================
doc.add_heading('七、关键文件索引', level=1)
files = [
    ('config.py', '全局配置，API Key、模型名、路径统一管理'),
    ('scripts/crawl_gb2760.py', 'GB2760-2024 数据爬虫'),
    ('scripts/process_data.py', '数据合并与去重处理'),
    ('scripts/build_vectorstore.py', '构建 ChromaDB 向量库'),
    ('src/models/llm.py', 'DeepSeek LLM 初始化（ChatOpenAI 兼容接口）'),
    ('src/models/embeddings.py', 'BGE-M3 Embedding 模型（ModelScope 下载）'),
    ('src/data/loader.py', 'JSON → LangChain Document 加载器'),
    ('src/data/vectorstore.py', 'ChromaDB 操作封装（检索、过滤）'),
    ('src/chains/interpret_chain.py', '配料解读 RAG Chain'),
    ('src/chains/risk_chain.py', '风险标注 RAG Chain'),
    ('src/agents/router.py', '意图路由器（关键词 + LLM fallback）'),
    ('src/agents/graph.py', 'LangGraph StateGraph Agent 编排'),
    ('src/prompts/', 'Prompt 模板（Markdown 格式，方便迭代）'),
    ('app/app.py', 'Streamlit 前端界面'),
]
for filename, desc in files:
    p = doc.add_paragraph()
    run = p.add_run(f'{filename}')
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    p.add_run(f' — {desc}')

doc.add_paragraph('')
doc.add_paragraph('— 文档结束 —').alignment = WD_ALIGN_PARAGRAPH.CENTER

# 保存
output_path = r'd:\食品配料智能分析agent\食鉴_FoodGuard_项目详解.docx'
doc.save(output_path)
print(f'Word 文档已生成: {output_path}')
